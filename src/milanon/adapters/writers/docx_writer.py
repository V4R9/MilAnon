"""DOCX writer — converts anonymized Markdown to Swiss Army DOCX format.

Uses the official CH Armee template (befehl_vorlage.docx) for style definitions.
Handles headings, bullet lists, numbered lists, tables, inline formatting,
blockquotes, horizontal rules, and HTML comment stripping.

Fixes: BUG-005 (bold), BUG-006 (---), BUG-007 (HTML comments),
       BUG-008 (blockquotes), BUG-009 (multi-col tables),
       BUG-010 (paragraph spacing), BUG-011 (Aufträge table).
"""

from __future__ import annotations

import contextlib
import re
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from milanon.domain.anonymizer import LEGEND_PATTERN
from milanon.domain.entities import AnonymizedDocument

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
    from docx.text.paragraph import Paragraph

# Template path relative to package structure.
# __file__ is src/milanon/adapters/writers/docx_writer.py → 5 parents = repo root.
_TEMPLATE_PATH = (
    Path(__file__).resolve().parent.parent.parent.parent.parent
    / "data"
    / "templates"
    / "docx"
    / "befehl_vorlage.docx"
)

# Regex patterns for line classification.
_RE_HEADING = re.compile(r"^(#{1,6})\s+(.*)")
_RE_BULLET = re.compile(r"^[-*]\s+(.*)")
_RE_NUMBERED = re.compile(r"^\d+\.\s+(.*)")
_RE_TABLE_ROW = re.compile(r"^\|(.+)\|$")
_RE_TABLE_SEP = re.compile(r"^\|[\s:|-]+\|$")
_RE_HR = re.compile(r"^-{3,}\s*$")
_RE_BLOCKQUOTE = re.compile(r"^>\s?(.*)")
_RE_HTML_COMMENT = re.compile(r"<!--.*?-->", re.DOTALL)

# Inline formatting patterns — order matters (longest match first).
_RE_INLINE = re.compile(
    r"(\*{3})(.+?)\1"  # ***bold+italic***
    r"|(\*{2})(.+?)\3"  # **bold**
    r"|(\*)(.+?)\5"  # *italic*
    r"|(`+)(.+?)\7"  # `code`
)

# Style mapping: heading level → DOCX style name.
_HEADING_STYLES: dict[int, str] = {
    1: "Heading 1",
    2: "Subject heading",
    3: "1. Main title",
    4: "1.1 Title",
    5: "1.1.1 Title",
    6: "Text Indent",  # Fallback for h6 — rendered bold.
}


def _parse_inline(text: str) -> list[tuple[str, bool, bool, bool]]:
    """Parse inline Markdown into (text, bold, italic, code) tuples.

    Handles ***bold+italic***, **bold**, *italic*, and `code`.
    Returns segments preserving the original text order.
    """
    segments: list[tuple[str, bool, bool, bool]] = []
    last_end = 0

    for m in _RE_INLINE.finditer(text):
        # Add any plain text before this match.
        if m.start() > last_end:
            segments.append((text[last_end : m.start()], False, False, False))

        if m.group(1):  # ***bold+italic***
            segments.append((m.group(2), True, True, False))
        elif m.group(3):  # **bold**
            segments.append((m.group(4), True, False, False))
        elif m.group(5):  # *italic*
            segments.append((m.group(6), False, True, False))
        elif m.group(7):  # `code`
            segments.append((m.group(8), False, False, True))

        last_end = m.end()

    # Trailing plain text.
    if last_end < len(text):
        segments.append((text[last_end:], False, False, False))

    return segments if segments else [(text, False, False, False)]


def _add_runs(paragraph: Paragraph, text: str) -> None:
    """Add inline-formatted runs to a paragraph."""
    for content, bold, italic, code in _parse_inline(text):
        if not content:
            continue
        run = paragraph.add_run(content)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if code:
            run.font.name = "Courier New"
            run.font.size = Pt(9)


def _load_template() -> DocxDocument:
    """Load the CH Armee template, clearing placeholder content.

    Falls back to an empty Document if the template is not found.
    """
    if _TEMPLATE_PATH.exists():
        doc = Document(str(_TEMPLATE_PATH))
        # Clear all existing paragraphs (keep style definitions).
        for p in list(doc.paragraphs):
            p_element = p._element
            p_element.getparent().remove(p_element)
        # Clear all existing tables.
        for t in list(doc.tables):
            t_element = t._element
            t_element.getparent().remove(t_element)
        return doc
    return Document()


def _add_page_break(doc: DocxDocument) -> None:
    """Insert a page break into the document."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _emit_table(doc: DocxDocument, rows: list[list[str]]) -> None:
    """Create a DOCX table from parsed Markdown table rows.

    Args:
        doc: The target DOCX document.
        rows: List of rows, each row a list of cell strings.
              rows[0] is the header row, rows[1:] are data rows.
    """
    if not rows:
        return

    n_cols = max(len(r) for r in rows)
    n_rows = len(rows)

    table = doc.add_table(rows=n_rows, cols=n_cols)

    # Apply Table Grid style if available.
    with contextlib.suppress(KeyError):
        table.style = "Table Grid"

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < n_cols:
                cell = table.rows[row_idx].cells[col_idx]
                # Clear default empty paragraph, add formatted content.
                cell.text = ""
                para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                _add_runs(para, cell_text.strip())

                # Bold the header row.
                if row_idx == 0:
                    for run in para.runs:
                        run.bold = True


def _strip_html_comments(text: str) -> str:
    """Remove all HTML comments from text."""
    return _RE_HTML_COMMENT.sub("", text)


class DocxWriter:
    """Writes anonymized Markdown content to a Swiss Army DOCX document.

    Loads the official CH Armee template for style definitions and maps
    Markdown elements to the correct DOCX paragraph/table styles.
    """

    def write(self, document: AnonymizedDocument, output_path: Path) -> Path:
        """Write anonymized content to a DOCX file.

        Args:
            document: The anonymized document.
            output_path: Destination .docx path (created if needed).

        Returns:
            The path to the written file.
        """
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Strip legend and HTML comments.
        text = LEGEND_PATTERN.sub("", document.content).strip()
        text = _strip_html_comments(text)

        doc = _load_template()

        # State for table accumulation.
        table_rows: list[list[str]] = []

        for line in text.splitlines():
            # Flush accumulated table rows if this line is not a table line.
            is_table_line = _RE_TABLE_ROW.match(line) or _RE_TABLE_SEP.match(line)
            if not is_table_line and table_rows:
                _emit_table(doc, table_rows)
                table_rows = []

            # --- Table rows ---
            if _RE_TABLE_SEP.match(line):
                # Separator line — skip.
                continue

            m_table = _RE_TABLE_ROW.match(line)
            if m_table:
                cells = [c.strip() for c in m_table.group(1).split("|")]
                table_rows.append(cells)
                continue

            # --- Empty lines → skip (BUG-010 fix) ---
            if not line.strip():
                continue

            # --- Horizontal rules → page break (BUG-006 fix) ---
            if _RE_HR.match(line):
                _add_page_break(doc)
                continue

            # --- Headings ---
            m_heading = _RE_HEADING.match(line)
            if m_heading:
                level = len(m_heading.group(1))
                heading_text = m_heading.group(2).strip()
                style_name = _HEADING_STYLES.get(level, "Text Indent")

                try:
                    para = doc.add_paragraph(style=style_name)
                except KeyError:
                    # Style not in template — fall back.
                    para = doc.add_paragraph()

                _add_runs(para, heading_text)

                # h6 fallback: make bold.
                if level == 6:
                    for run in para.runs:
                        run.bold = True
                continue

            # --- Bullet list ---
            m_bullet = _RE_BULLET.match(line)
            if m_bullet:
                try:
                    para = doc.add_paragraph(style="Bullet List 1")
                except KeyError:
                    para = doc.add_paragraph()
                _add_runs(para, m_bullet.group(1))
                continue

            # --- Numbered list ---
            m_numbered = _RE_NUMBERED.match(line)
            if m_numbered:
                try:
                    para = doc.add_paragraph(style="Numbered List 2")
                except KeyError:
                    para = doc.add_paragraph()
                _add_runs(para, m_numbered.group(1))
                continue

            # --- Blockquote (BUG-008 fix) ---
            m_blockquote = _RE_BLOCKQUOTE.match(line)
            if m_blockquote:
                try:
                    para = doc.add_paragraph(style="Text Indent")
                except KeyError:
                    para = doc.add_paragraph()
                quote_text = m_blockquote.group(1)
                run = para.add_run(quote_text)
                run.italic = True
                continue

            # --- Normal text → Text Indent ---
            try:
                para = doc.add_paragraph(style="Text Indent")
            except KeyError:
                para = doc.add_paragraph()
            _add_runs(para, line)

        # Flush any remaining table rows.
        if table_rows:
            _emit_table(doc, table_rows)

        doc.save(str(output_path))
        return output_path

    def default_extension(self) -> str:
        return ".docx"
