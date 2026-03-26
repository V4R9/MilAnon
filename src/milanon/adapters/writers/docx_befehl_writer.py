"""DOCX Befehl writer — converts structured Markdown into a formatted DOCX
using the official Armee-Vorlage as template.

Style mapping (verified against befehl_vorlage.docx — all styles confirmed present):
  # Grundlagen       → "Heading 1"          ✓ exists in template
  ## "DECKNAME"      → "Subject heading"     ✓ exists in template
  ### 1 Title        → "1. Main title"       ✓ exists in template
  #### 1.1 Title     → "1.1 Title"           ✓ exists in template
  ##### 1.1.1 Title  → "1.1.1 Title"         ✓ exists in template
  Body text          → "Text Indent"         ✓ exists in template
  - Bullet item      → "Bullet List 1"       ✓ exists in template
  1. Item            → "Numbered List 2"     ✓ exists in template
  | Table |          → DOCX Table (Aufträge) — table cells use "Bullet List 1"

BUG-019: Numbered heading detection (1. / 1.1. / 1.2.1. patterns from PDF extraction)
BUG-020: PDF artifact stripping (page footers, .docx filenames, page numbers)
BUG-021: Inline formatting (bold/italic) via _parse_inline / _add_runs
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from milanon.domain.anonymizer import LEGEND_PATTERN

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph

# ---------------------------------------------------------------------------
# Regex patterns
# ---------------------------------------------------------------------------

# Markdown table rows (pipe-delimited)
_TABLE_ROW_RE = re.compile(r"^\|(.+)\|$")
# Separator row in Markdown tables (e.g. |---|---|)
_TABLE_SEP_RE = re.compile(r"^\|[\s\-:|]+\|$")
# Ordered list item
_ORDERED_LIST_RE = re.compile(r"^\d+\.\s+(.*)$")

# BUG-019: Numbered heading patterns (checked BEFORE _ORDERED_LIST_RE)
# Sub-subsection must be checked first (most specific).
_RE_HEADING_SUB_SUB = re.compile(r"^(\d\.\d+\.\d+\.?\s+.*)$")
# Subsection
_RE_HEADING_SUB = re.compile(r"^(\d\.\d+\.?\s+.*)$")
# Main section: single digit followed by capital letter (to avoid list items)
_RE_HEADING_MAIN = re.compile(r"^(\d\.\s+[A-ZÄÖÜ].*)$")

# Horizontal rule
_RE_HR = re.compile(r"^-{3,}\s*$")
# Blockquote
_RE_BLOCKQUOTE = re.compile(r"^>\s?(.*)")
# HTML comments
_RE_HTML_COMMENT = re.compile(r"<!--.*?-->", re.DOTALL)

# Inline formatting patterns — order matters (longest match first).
_RE_INLINE = re.compile(
    r"(\*{3})(.+?)\1"  # ***bold+italic***
    r"|(\*{2})(.+?)\3"  # **bold**
    r"|(\*)(.+?)\5"  # *italic*
    r"|(`+)(.+?)\7"  # `code`
)

# BUG-020: PDF artifact patterns
_RE_ARTIFACT_DOCX = re.compile(r"^.*\.docx\s*$", re.IGNORECASE)
_RE_ARTIFACT_PAGE_NUM = re.compile(r"^\d+/\d+\s*$")
_RE_ARTIFACT_FOOTER = re.compile(r'^.+".+".+\d+/\d+\s*$')


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------


def _parse_inline(text: str) -> list[tuple[str, bool, bool, bool]]:
    """Parse inline Markdown into (text, bold, italic, code) tuples.

    Handles ***bold+italic***, **bold**, *italic*, and `code`.
    Returns segments preserving the original text order.
    """
    segments: list[tuple[str, bool, bool, bool]] = []
    last_end = 0

    for m in _RE_INLINE.finditer(text):
        if m.start() > last_end:
            segments.append((text[last_end : m.start()], False, False, False))

        if m.group(1):  # ***bold+italic***
            segments.append((m.group(2), True, True, False))
        elif m.group(3):  # **bold**
            segments.append((m.group(4), True, False, False))
        elif m.group(5):  # *italic*
            segments.append((m.group(6), False, True, False))
        elif m.group(7):  # `code`
            segments.append((m.group(8), False, False, True))

        last_end = m.end()

    if last_end < len(text):
        segments.append((text[last_end:], False, False, False))

    return segments if segments else [(text, False, False, False)]


def _add_runs(paragraph: Paragraph, text: str) -> None:
    """Add inline-formatted runs to a paragraph."""
    for content, bold, italic, code in _parse_inline(text):
        if not content:
            continue
        run = paragraph.add_run(content)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if code:
            run.font.name = "Courier New"
            run.font.size = Pt(9)


def _strip_html_comments(text: str) -> str:
    """Remove all HTML comments from text."""
    return _RE_HTML_COMMENT.sub("", text)


def _add_page_break(doc: Document) -> None:
    """Insert a page break into the document."""
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)  # noqa: SLF001


def _strip_pdf_artifacts(lines: list[str]) -> list[str]:
    """Remove PDF extraction artifacts (page footers, filenames, page numbers).

    BUG-020: Filters out:
    - Lines ending with .docx (source filenames)
    - Lines that are only a page number like 1/7
    - Footer lines containing a quoted string and ending with N/M page number
    """
    result = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            result.append(line)
            continue
        if _RE_ARTIFACT_DOCX.match(stripped):
            continue
        if _RE_ARTIFACT_PAGE_NUM.match(stripped):
            continue
        if _RE_ARTIFACT_FOOTER.match(stripped):
            continue
        result.append(line)
    return result


# ---------------------------------------------------------------------------
# Writer class
# ---------------------------------------------------------------------------


class DocxBefehlWriter:
    """Converts structured Markdown into a DOCX using the Armee-Vorlage template."""

    # Column widths for Aufträge table (Element | Auftrag)
    COL_WIDTH_LEFT = Cm(4.0)
    COL_WIDTH_RIGHT = Cm(4.3)

    def write(
        self, markdown_text: str, template_path: Path, output_path: Path
    ) -> Path:
        """Convert Markdown text to a DOCX file based on the template.

        Args:
            markdown_text: Structured Markdown content (Befehl).
            template_path: Path to befehl_vorlage.docx.
            output_path: Where to save the generated DOCX.

        Returns:
            The path to the written DOCX file.
        """
        doc = Document(str(template_path))
        self._clear_body(doc)

        # Pre-processing: strip legend, HTML comments, PDF artifacts
        text = LEGEND_PATTERN.sub("", markdown_text).strip()
        text = _strip_html_comments(text)

        lines = text.splitlines()
        lines = _strip_pdf_artifacts(lines)

        i = 0
        while i < len(lines):
            line = lines[i]

            # Detect Markdown table block
            if _TABLE_ROW_RE.match(line):
                table_lines = []
                while i < len(lines) and _TABLE_ROW_RE.match(lines[i]):
                    if not _TABLE_SEP_RE.match(lines[i]):
                        table_lines.append(lines[i])
                    i += 1
                # Skip header row (first line is column headers)
                data_rows = (
                    table_lines[1:] if len(table_lines) > 1 else table_lines
                )
                self._add_table(doc, data_rows)
                continue

            self._add_line(doc, line)
            i += 1

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path

    def _clear_body(self, doc: Document) -> None:
        """Remove all paragraphs and tables from the document body."""
        body = doc.element.body
        for child in list(body):
            tag = (
                child.tag.split("}")[-1] if "}" in child.tag else child.tag
            )
            if tag in ("p", "tbl"):
                body.remove(child)

    def _add_line(self, doc: Document, line: str) -> None:
        """Map a single Markdown line to a styled DOCX paragraph."""
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("", style="Text Indent")
            return

        # --- Horizontal rule → page break ---
        if _RE_HR.match(stripped):
            _add_page_break(doc)
            return

        # --- Markdown # headings (order: ##### before ####, etc.) ---
        if stripped.startswith("##### "):
            text = stripped[6:]
            para = doc.add_paragraph(style="1.1.1 Title")
            _add_runs(para, text)
        elif stripped.startswith("#### "):
            text = stripped[5:]
            para = doc.add_paragraph(style="1.1 Title")
            _add_runs(para, text)
        elif stripped.startswith("### "):
            text = stripped[4:]
            para = doc.add_paragraph(style="1. Main title")
            _add_runs(para, text)
        elif stripped.startswith("## "):
            text = stripped[3:]
            para = doc.add_paragraph(style="Subject heading")
            _add_runs(para, text)
        elif stripped.startswith("# "):
            text = stripped[2:]
            para = doc.add_paragraph(style="Heading 1")
            _add_runs(para, text)

        # --- BUG-019: Numbered heading patterns (before ordered list!) ---
        elif _RE_HEADING_SUB_SUB.match(stripped):
            text = _RE_HEADING_SUB_SUB.match(stripped).group(1)
            para = doc.add_paragraph(style="1.1.1 Title")
            _add_runs(para, text)
        elif _RE_HEADING_SUB.match(stripped):
            text = _RE_HEADING_SUB.match(stripped).group(1)
            para = doc.add_paragraph(style="1.1 Title")
            _add_runs(para, text)
        elif _RE_HEADING_MAIN.match(stripped):
            text = _RE_HEADING_MAIN.match(stripped).group(1)
            para = doc.add_paragraph(style="1. Main title")
            _add_runs(para, text)

        # --- Bullet list ---
        elif stripped.startswith("- "):
            text = stripped[2:]
            para = doc.add_paragraph(style="Bullet List 1")
            _add_runs(para, text)

        # --- Ordered list (only if NOT caught by heading patterns above) ---
        elif (m := _ORDERED_LIST_RE.match(stripped)):
            text = m.group(1)
            para = doc.add_paragraph(style="Numbered List 2")
            _add_runs(para, text)

        # --- Blockquote ---
        elif (m_bq := _RE_BLOCKQUOTE.match(stripped)):
            para = doc.add_paragraph(style="Text Indent")
            run = para.add_run(m_bq.group(1))
            run.italic = True

        # --- Normal text ---
        else:
            para = doc.add_paragraph(style="Text Indent")
            _add_runs(para, stripped)

    def _add_table(self, doc: Document, data_rows: list[str]) -> None:
        """Create an Auftraege table (Nx2) from Markdown table rows."""
        if not data_rows:
            return

        table = doc.add_table(rows=len(data_rows), cols=2)

        for row in table.rows:
            row.cells[0].width = self.COL_WIDTH_LEFT
            row.cells[1].width = self.COL_WIDTH_RIGHT

        for row_idx, row_text in enumerate(data_rows):
            cells = [c.strip() for c in row_text.strip("|").split("|")]
            if len(cells) < 2:
                continue

            element_name = cells[0].strip()
            auftrag_text = cells[1].strip()

            # Left column: Einheitsbezeichnung
            left_cell = table.rows[row_idx].cells[0]
            left_cell.text = ""
            left_para = (
                left_cell.paragraphs[0]
                if left_cell.paragraphs
                else left_cell.add_paragraph()
            )
            _add_runs(left_para, element_name)

            # Right column: Auftrag lines with Bullet List 1 style
            right_cell = table.rows[row_idx].cells[1]
            right_cell.text = ""

            auftrag_lines = [
                al.strip().lstrip("- ").strip()
                for al in re.split(r"<br>|;\s*(?=-)", auftrag_text)
                if al.strip()
            ]
            if not auftrag_lines:
                auftrag_lines = [auftrag_text]

            for line_idx, aline in enumerate(auftrag_lines):
                p = (
                    right_cell.paragraphs[0]
                    if line_idx == 0
                    else right_cell.add_paragraph()
                )
                p.style = doc.styles["Bullet List 1"]
                _add_runs(p, aline)
