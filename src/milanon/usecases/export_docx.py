"""ExportDocxUseCase — export a Markdown Befehl to DOCX with optional de-anonymization."""

from __future__ import annotations

import re
from pathlib import Path

# TODO(FR-004): Move python-docx usage to adapter layer (DocxWriter)
# Currently violates Dependency Rule — UseCase imports infra library directly
from docx import Document

from milanon.domain.anonymizer import PLACEHOLDER_PATTERN


class ExportDocxUseCase:
    """Export a Markdown Befehl to DOCX with optional de-anonymization."""

    def __init__(self, repository, writer):
        self._repo = repository
        self._writer = writer  # DocxBefehlWriter

    def execute(
        self,
        input_path: Path,
        output_path: Path,
        template_path: Path,
        deanonymize: bool = False,
    ) -> Path:
        """Convert Markdown to DOCX, optionally replacing placeholders with cleartext.

        Args:
            input_path: Path to the Markdown file.
            output_path: Path for the output .docx file.
            template_path: Path to befehl_vorlage.docx.
            deanonymize: If True, replace [TYPE_NNN] placeholders with cleartext.

        Returns:
            The path to the generated DOCX.
        """
        markdown_text = input_path.read_text(encoding="utf-8")

        # Step 1: Optionally de-anonymize the Markdown text before conversion
        if deanonymize:
            markdown_text = self._deanonymize_text(markdown_text)

        # Step 2: Convert Markdown → DOCX
        result_path = self._writer.write(markdown_text, template_path, output_path)

        # Step 3: If de-anonymizing, also replace any remaining placeholders
        # in the generated DOCX (tables, headers, etc.)
        if deanonymize:
            doc = Document(str(result_path))
            self._deanonymize_docx(doc)
            doc.save(str(result_path))

        return result_path

    def _deanonymize_text(self, text: str) -> str:
        """Replace all placeholders in plain text with cleartext values."""
        def _replace(match: re.Match) -> str:
            placeholder = match.group(0)
            mapping = self._repo.get_placeholder(placeholder)
            if mapping:
                return mapping.original_value
            return placeholder

        return PLACEHOLDER_PATTERN.sub(_replace, text)

    def _deanonymize_docx(self, doc: Document) -> None:
        """Replace all placeholders in the DOCX paragraphs and table cells."""
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph)
        # Also handle headers and footers
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                self._replace_in_paragraph(paragraph)
            for paragraph in section.footer.paragraphs:
                self._replace_in_paragraph(paragraph)

    def _replace_in_paragraph(self, paragraph) -> None:
        """Replace placeholders within a paragraph's runs, preserving formatting."""
        full_text = paragraph.text
        if not PLACEHOLDER_PATTERN.search(full_text):
            return

        new_text = PLACEHOLDER_PATTERN.sub(self._resolve_placeholder, full_text)
        if new_text == full_text:
            return

        # Rewrite runs: put all text in first run, clear the rest
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.text = new_text

    def _resolve_placeholder(self, match: re.Match) -> str:
        """Resolve a single placeholder match to cleartext."""
        placeholder = match.group(0)
        mapping = self._repo.get_placeholder(placeholder)
        if mapping:
            return mapping.original_value
        return placeholder
