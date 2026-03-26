"""Comprehensive test suite for FR-004 DOCX Writer rewrite.

Tests verify that the new DocxWriter correctly converts Markdown to DOCX
with proper styles, inline formatting, tables, and stripping of artifacts.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from milanon.adapters.writers.docx_writer import DocxWriter
from milanon.domain.entities import AnonymizedDocument, DocumentFormat

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_doc(content: str) -> AnonymizedDocument:
    """Create a minimal AnonymizedDocument for testing."""
    return AnonymizedDocument(
        source_path="test.md",
        output_format=DocumentFormat.MARKDOWN,
        content=content,
    )


def _write_and_open(content: str, tmp_path: Path) -> Document:
    """Write content through DocxWriter and return the opened DOCX Document."""
    writer = DocxWriter()
    out = tmp_path / "output.docx"
    writer.write(_make_doc(content), out)
    return Document(str(out))


def _paragraphs_with_style(docx_doc: Document, style_name: str) -> list:
    """Return all paragraphs matching a given style name."""
    return [p for p in docx_doc.paragraphs if p.style.name == style_name]


def _all_text(docx_doc: Document) -> str:
    """Return the full text of all paragraphs concatenated."""
    return "\n".join(p.text for p in docx_doc.paragraphs)


# ===========================================================================
# Heading Styles
# ===========================================================================

class TestHeadingStyles:
    """Tests 1-5: Markdown headings map to correct DOCX styles."""

    def test_heading1_uses_heading1_style(self, tmp_path: Path) -> None:
        docx_doc = _write_and_open("# Einsatzbefehl", tmp_path)
        matches = _paragraphs_with_style(docx_doc, "Heading 1")
        assert len(matches) >= 1
        assert "Einsatzbefehl" in matches[0].text

    def test_heading2_uses_subject_heading_style(self, tmp_path: Path) -> None:
        docx_doc = _write_and_open("## Lage", tmp_path)
        matches = _paragraphs_with_style(docx_doc, "Subject heading")
        assert len(matches) >= 1
        assert "Lage" in matches[0].text

    def test_heading3_uses_main_title_style(self, tmp_path: Path) -> None:
        docx_doc = _write_and_open("### Allgemeine Lage", tmp_path)
        matches = _paragraphs_with_style(docx_doc, "1. Main title")
        assert len(matches) >= 1
        assert "Allgemeine Lage" in matches[0].text

    def test_heading4_uses_title_style(self, tmp_path: Path) -> None:
        docx_doc = _write_and_open("#### Feind", tmp_path)
        matches = _paragraphs_with_style(docx_doc, "1.1 Title")
        assert len(matches) >= 1
        assert "Feind" in matches[0].text

    def test_heading5_uses_sub_title_style(self, tmp_path: Path) -> None:
        docx_doc = _write_and_open("##### Details", tmp_path)
        matches = _paragraphs_with_style(docx_doc, "1.1.1 Title")
        assert len(matches) >= 1
        assert "Details" in matches[0].text


# ===========================================================================
# Inline Formatting (BUG-005)
# ===========================================================================

class TestInlineFormatting:
    """Tests 6-9: Bold, italic, inline code are rendered correctly."""

    def test_bold_text_creates_bold_run(self, tmp_path: Path) -> None:
        docx_doc = _write_and_open("This is **bold** text", tmp_path)
        # Find the paragraph containing our text
        para = next(p for p in docx_doc.paragraphs if "bold" in p.text)
        bold_runs = [r for r in para.runs if r.bold]
        assert len(bold_runs) >= 1
        assert any("bold" in r.text for r in bold_runs)
        # The literal ** markers should not appear
        assert "**" not in para.text

    def test_italic_text_creates_italic_run(self, tmp_path: Path) -> None:
        docx_doc = _write_and_open("This is *italic* text", tmp_path)
        para = next(p for p in docx_doc.paragraphs if "italic" in p.text)
        italic_runs = [r for r in para.runs if r.italic]
        assert len(italic_runs) >= 1
        assert any("italic" in r.text for r in italic_runs)
        assert para.text.count("*") == 0

    def test_mixed_inline_formatting(self, tmp_path: Path) -> None:
        docx_doc = _write_and_open("Normal **bold** and *italic* end", tmp_path)
        para = next(p for p in docx_doc.paragraphs if "Normal" in p.text)
        # Should have multiple runs for the different formatting segments
        assert len(para.runs) >= 4
        bold_runs = [r for r in para.runs if r.bold]
        italic_runs = [r for r in para.runs if r.italic]
        assert len(bold_runs) >= 1
        assert len(italic_runs) >= 1
        # No literal markdown markers
        assert "**" not in para.text
        full_text = para.text
        assert "bold" in full_text
        assert "italic" in full_text

    def test_inline_code_uses_monospace(self, tmp_path: Path) -> None:
        docx_doc = _write_and_open("Use `milanon pack` command", tmp_path)
        para = next(p for p in docx_doc.paragraphs if "milanon" in p.text)
        # Find the run containing the code text
        code_runs = [
            r for r in para.runs
            if "milanon" in r.text and r.font.name is not None
        ]
        # The code run should use a monospace font
        assert len(code_runs) >= 1
        mono_fonts = {"Courier New", "Courier", "Consolas", "Monaco", "monospace"}
        assert any(r.font.name in mono_fonts for r in code_runs)
        # Backtick markers should not appear
        assert "`" not in para.text


# ===========================================================================
# Bullet Lists
# ===========================================================================

class TestBulletLists:
    """Tests 10-11: List items use correct DOCX styles."""

    def test_bullet_list_uses_bullet_style(self, tmp_path: Path) -> None:
        content = "- Erste Aufgabe\n- Zweite Aufgabe\n- Dritte Aufgabe"
        docx_doc = _write_and_open(content, tmp_path)
        bullet_paras = _paragraphs_with_style(docx_doc, "Bullet List 1")
        assert len(bullet_paras) >= 3
        assert "Erste Aufgabe" in bullet_paras[0].text

    def test_numbered_list_uses_numbered_style(self, tmp_path: Path) -> None:
        content = "1. Erster Punkt\n2. Zweiter Punkt\n3. Dritter Punkt"
        docx_doc = _write_and_open(content, tmp_path)
        # Accept any style containing "Numbered" or "List Number"
        numbered_paras = [
            p for p in docx_doc.paragraphs
            if "Number" in p.style.name or "Numbered" in p.style.name
        ]
        assert len(numbered_paras) >= 3


# ===========================================================================
# Tables (BUG-009/011)
# ===========================================================================

class TestTables:
    """Tests 12-15: Markdown tables render as DOCX tables."""

    TWO_COL_TABLE = (
        "| Einheit | Auftrag |\n"
        "|---|---|\n"
        "| [EINHEIT_001] | Sichern |\n"
        "| [EINHEIT_002] | Aufklaeren |"
    )

    FOUR_COL_TABLE = (
        "| Akteur | Funktion | Standort | Bemerkung |\n"
        "|---|---|---|---|\n"
        "| [PERSON_001] | Kdt | [ORT_001] | Verantwortlich |\n"
        "| [PERSON_002] | Stv | [ORT_002] | Stellvertreter |"
    )

    def test_two_column_table(self, tmp_path: Path) -> None:
        docx_doc = _write_and_open(self.TWO_COL_TABLE, tmp_path)
        assert len(docx_doc.tables) >= 1
        table = docx_doc.tables[0]
        # 2 columns
        assert len(table.columns) == 2

    def test_four_column_table(self, tmp_path: Path) -> None:
        docx_doc = _write_and_open(self.FOUR_COL_TABLE, tmp_path)
        assert len(docx_doc.tables) >= 1
        table = docx_doc.tables[0]
        assert len(table.columns) == 4

    def test_table_header_row(self, tmp_path: Path) -> None:
        docx_doc = _write_and_open(self.TWO_COL_TABLE, tmp_path)
        table = docx_doc.tables[0]
        # First row should contain the header text
        header_texts = [cell.text.strip() for cell in table.rows[0].cells]
        assert "Einheit" in header_texts
        assert "Auftrag" in header_texts

    def test_table_separator_not_visible(self, tmp_path: Path) -> None:
        docx_doc = _write_and_open(self.TWO_COL_TABLE, tmp_path)
        full_text = _all_text(docx_doc)
        # Also check table cell text
        table_text = ""
        for table in docx_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text + "\n"
        combined = full_text + table_text
        # The separator line should not appear as literal text
        assert "|---|" not in combined
        assert "---" not in combined or combined.count("---") == 0


# ===========================================================================
# HTML Comments (BUG-007)
# ===========================================================================

class TestHtmlComments:
    """Tests 16-18: HTML comments are stripped from output."""

    def test_html_comments_stripped(self, tmp_path: Path) -> None:
        content = "Before <!-- FILL: Zeitraum --> After"
        docx_doc = _write_and_open(content, tmp_path)
        full_text = _all_text(docx_doc)
        assert "<!-- FILL:" not in full_text
        assert "FILL:" not in full_text

    def test_kdt_entscheid_comment_stripped(self, tmp_path: Path) -> None:
        content = "Entscheid: <!-- KDT ENTSCHEID: Variante A oder B --> weiter"
        docx_doc = _write_and_open(content, tmp_path)
        full_text = _all_text(docx_doc)
        assert "<!-- KDT ENTSCHEID:" not in full_text
        assert "KDT ENTSCHEID:" not in full_text

    def test_line_with_only_comment_produces_no_paragraph(self, tmp_path: Path) -> None:
        content = "Paragraph one\n<!-- FILL: placeholder -->\nParagraph two"
        docx_doc = _write_and_open(content, tmp_path)
        texts = [p.text.strip() for p in docx_doc.paragraphs if p.text.strip()]
        # Should only have the two real paragraphs, not one for the comment
        assert "FILL" not in " ".join(texts)
        # The comment-only line should not create a visible paragraph
        non_empty = [t for t in texts if t]
        assert len(non_empty) == 2


# ===========================================================================
# Horizontal Rules (BUG-006)
# ===========================================================================

class TestHorizontalRules:
    """Test 19: Horizontal rules do not appear as literal text."""

    def test_horizontal_rule_not_rendered_as_text(self, tmp_path: Path) -> None:
        content = "Section one\n\n---\n\nSection two"
        docx_doc = _write_and_open(content, tmp_path)
        # The literal --- should not appear as paragraph text
        texts = [p.text.strip() for p in docx_doc.paragraphs]
        assert "---" not in texts


# ===========================================================================
# Blockquotes (BUG-008)
# ===========================================================================

class TestBlockquotes:
    """Test 20: Blockquotes render with italic formatting."""

    def test_blockquote_rendered_as_italic(self, tmp_path: Path) -> None:
        content = "> Dies ist ein Hinweis"
        docx_doc = _write_and_open(content, tmp_path)
        # Find the paragraph containing the blockquote text
        para = next(
            (p for p in docx_doc.paragraphs if "Hinweis" in p.text), None
        )
        assert para is not None, "Blockquote text not found in output"
        # The > marker should be stripped
        assert not para.text.startswith(">")
        # Text should be italic (either via style or run formatting)
        has_italic = any(r.italic for r in para.runs)
        is_italic_style = "quote" in para.style.name.lower() or "italic" in para.style.name.lower()
        assert has_italic or is_italic_style, (
            f"Blockquote not italic. Style: {para.style.name}, "
            f"runs italic: {[r.italic for r in para.runs]}"
        )


# ===========================================================================
# Paragraph Spacing (BUG-010)
# ===========================================================================

class TestParagraphSpacing:
    """Test 21: Empty lines do not produce empty paragraphs."""

    def test_empty_lines_not_rendered(self, tmp_path: Path) -> None:
        content = "First paragraph\n\n\n\nSecond paragraph"
        docx_doc = _write_and_open(content, tmp_path)
        empty_paras = [p for p in docx_doc.paragraphs if p.text.strip() == ""]
        # Allow at most 1 empty paragraph (e.g., for spacing), but not 3+
        assert len(empty_paras) <= 1, (
            f"Too many empty paragraphs: {len(empty_paras)}"
        )


# ===========================================================================
# Integration Tests
# ===========================================================================

class TestIntegration:
    """Tests 22-25: End-to-end integration tests."""

    EINSATZBEFEHL_MD = """\
# Einsatzbefehl [EINHEIT_001]

## 1 Lage

### 1.1 Allgemeine Lage

Die Bedrohungslage im Raum [ORT_001] hat sich verschaerft.

- Erhoehte Aktivitaet im Sektor Nord
- Aufklaerungsergebnisse bestaetigen Feindpraesenz
- **Wichtig**: Sofortmassnahmen eingeleitet

### 1.2 Eigene Lage

[EINHEIT_001] befindet sich im Bereitschaftsraum [ORT_002].

| Einheit | Staerke | Bereitschaft |
|---|---|---|
| [EINHEIT_002] | 120 | Hoch |
| [EINHEIT_003] | 95 | Mittel |

## 2 Auftrag

[EINHEIT_001] sichert den Raum [ORT_001] ab sofort.

> Hinweis: Koordination mit [EINHEIT_004] erforderlich.

## 3 Durchfuehrung

### 3.1 Absicht Kdt

**Absicht**: Sicherung des Raums [ORT_001] durch Praesenz.

<!-- KDT ENTSCHEID: Variante A gewaehlt -->

#### Auftraege

| Einheit | Auftrag |
|---|---|
| [EINHEIT_002] | Sichern Sektor Nord |
| [EINHEIT_003] | Reserve bereithalten |

---

## 4 Logistik

- Munition: Grunddotation
- Verpflegung: 48h autonom

## 5 Fuehrung

Kdt [PERSON_001] fuehrt ab KP [ORT_003].
"""

    def test_full_einsatzbefehl_structure(self, tmp_path: Path) -> None:
        docx_doc = _write_and_open(self.EINSATZBEFEHL_MD, tmp_path)

        # Headings present
        h1 = _paragraphs_with_style(docx_doc, "Heading 1")
        assert len(h1) >= 1, "Missing Heading 1"

        h2 = _paragraphs_with_style(docx_doc, "Subject heading")
        assert len(h2) >= 4, f"Expected >=4 Subject headings, got {len(h2)}"

        h3 = _paragraphs_with_style(docx_doc, "1. Main title")
        assert len(h3) >= 2, f"Expected >=2 Main titles, got {len(h3)}"

        # Tables present
        assert len(docx_doc.tables) >= 2, (
            f"Expected >=2 tables, got {len(docx_doc.tables)}"
        )

        # Bullet items present
        bullets = _paragraphs_with_style(docx_doc, "Bullet List 1")
        assert len(bullets) >= 2, f"Expected >=2 bullet items, got {len(bullets)}"

        # No HTML comments visible
        full_text = _all_text(docx_doc)
        assert "<!-- KDT ENTSCHEID" not in full_text
        assert "KDT ENTSCHEID" not in full_text

        # Bold formatting applied (not literal **)
        assert "**" not in full_text

        # Blockquote rendered (> stripped)
        assert full_text.count("> Hinweis") == 0
        assert "Hinweis" in full_text

    def test_legend_stripped(self, tmp_path: Path) -> None:
        content = (
            "# Test\n\n"
            "Some content with [PERSON_001].\n\n"
            "<!-- MILANON LEGEND START\n"
            "[PERSON_001]: PERSON\n"
            "[ORT_001]: ORT\n"
            "MILANON LEGEND END -->"
        )
        docx_doc = _write_and_open(content, tmp_path)
        full_text = _all_text(docx_doc)
        assert "MILANON LEGEND" not in full_text
        assert "[PERSON_001]: PERSON" not in full_text
        # But the placeholder in the content should remain
        assert "[PERSON_001]" in full_text

    def test_output_is_valid_docx(self, tmp_path: Path) -> None:
        writer = DocxWriter()
        out = tmp_path / "valid.docx"
        writer.write(_make_doc("# Test\n\nHello world"), out)
        # Should not raise
        doc = Document(str(out))
        assert len(doc.paragraphs) >= 1

    def test_writer_returns_output_path(self, tmp_path: Path) -> None:
        writer = DocxWriter()
        out = tmp_path / "result.docx"
        result = writer.write(_make_doc("# Test"), out)
        assert result == out
        assert out.exists()
