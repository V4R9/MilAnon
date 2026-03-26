"""Fixtures for parser adapter tests."""

from __future__ import annotations

import base64
import io
import struct
import zlib
from pathlib import Path

import pytest

from milanon.adapters.parsers.docx_parser import DocxParser
from milanon.adapters.parsers.eml_parser import EmlParser
from milanon.adapters.parsers.pdf_parser import PdfParser, _tesseract_available
from milanon.adapters.parsers.xlsx_csv_parser import XlsxCsvParser

# True if Tesseract binary is installed on this machine
TESSERACT_AVAILABLE = _tesseract_available()

FIXTURES_DIR = Path(__file__).parent.parent.parent / "e2e" / "fixtures"

# Known body text for the base64 EML fixture
SAMPLE_BASE64_BODY = "Auftrag erhalten. Dossier ist bereit."

# Known content in the simple DOCX fixture
SIMPLE_DOCX_TEXTS = ["Befehl Inf Kp 56/1", "Hptm Marco BERNASCONI", "756.1234.5678.97"]


# ---------------------------------------------------------------------------
# EML fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def eml_parser() -> EmlParser:
    return EmlParser()


@pytest.fixture
def sample_eml_path() -> Path:
    return FIXTURES_DIR / "sample.eml"


@pytest.fixture
def sample_multipart_eml_path() -> Path:
    return FIXTURES_DIR / "sample_multipart.eml"


@pytest.fixture
def sample_signature_eml_path() -> Path:
    return FIXTURES_DIR / "sample_signature.eml"


@pytest.fixture
def sample_attachment_eml_path() -> Path:
    return FIXTURES_DIR / "sample_attachment.eml"


@pytest.fixture
def sample_base64_eml_path(tmp_path: Path) -> Path:
    """Generate a valid base64-encoded EML in tmp_path using Python's base64 module."""
    body_b64 = base64.b64encode(SAMPLE_BASE64_BODY.encode("utf-8")).decode("ascii")
    content = (
        "From: sender@example.com\r\n"
        "To: recipient@example.com\r\n"
        "Subject: Base64 Test\r\n"
        "Date: Mon, 23 Mar 2026 09:00:00 +0100\r\n"
        "MIME-Version: 1.0\r\n"
        "Content-Type: text/plain; charset=utf-8\r\n"
        "Content-Transfer-Encoding: base64\r\n"
        "\r\n"
        f"{body_b64}\r\n"
    )
    path = tmp_path / "sample_base64.eml"
    path.write_text(content, encoding="ascii")
    return path


# ---------------------------------------------------------------------------
# DOCX fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def docx_parser() -> DocxParser:
    return DocxParser()


@pytest.fixture
def simple_docx_path(tmp_path: Path) -> Path:
    """DOCX with basic paragraphs — content matches SIMPLE_DOCX_TEXTS."""
    from docx import Document

    doc = Document()
    doc.add_heading("Befehl Inf Kp 56/1", 0)
    doc.add_paragraph("Kommandant: Hptm Marco BERNASCONI")
    doc.add_paragraph("AHV-Nr: 756.1234.5678.97")
    doc.add_paragraph("Standort: Kaserne Wangen an der Aare")
    path = tmp_path / "simple.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def docx_with_table_path(tmp_path: Path) -> Path:
    """DOCX with a 3-column personnel table preceded by a paragraph."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Personalliste Inf Kp 56/1")
    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Grad"
    table.cell(0, 2).text = "Telefon"
    table.cell(1, 0).text = "Hans Muster"
    table.cell(1, 1).text = "Hptm"
    table.cell(1, 2).text = "079 535 80 46"
    table.cell(2, 0).text = "Peter Meier"
    table.cell(2, 1).text = "Lt"
    table.cell(2, 2).text = "079 123 45 67"
    path = tmp_path / "with_table.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def docx_with_header_footer_path(tmp_path: Path) -> Path:
    """DOCX with a header and a footer containing PII-rich text."""
    from docx import Document

    doc = Document()
    section = doc.sections[0]
    section.header.paragraphs[0].text = "VERTRAULICH - Inf Kp 56/1"
    section.footer.paragraphs[0].text = "Hans Muster | hans.muster@army.ch"
    doc.add_paragraph("Hauptinhalt des Dokuments")
    path = tmp_path / "with_header_footer.docx"
    doc.save(str(path))
    return path


@pytest.fixture
def docx_with_images_path(tmp_path: Path) -> Path:
    """DOCX with 2 embedded inline images surrounding body text."""
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_paragraph("Dokument mit eingebetteten Bildern")
    png_data = _make_minimal_png()
    doc.add_picture(io.BytesIO(png_data), width=Inches(0.5))
    doc.add_picture(io.BytesIO(png_data), width=Inches(0.5))
    doc.add_paragraph("Text nach den Bildern")
    path = tmp_path / "with_images.docx"
    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# PDF fixtures (fpdf2 + Pillow)
# ---------------------------------------------------------------------------

# Known text used in the simple PDF fixture (must produce >= 50 chars after extraction)
SIMPLE_PDF_TEXT = (
    "Kommandant: Hptm Marco BERNASCONI\nAHV-Nr: 756.1234.5678.97\nEinheit: Inf Kp 56/1"
)

# Text used on each page of the multi-page PDF fixture
MULTI_PAGE_PDF_TEXTS = [
    "Seite 1: Hans Muster, hans.muster@army.ch",
    "Seite 2: AHV 756.9876.5432.10, Tel 079 123 45 67",
]


@pytest.fixture
def pdf_parser() -> PdfParser:
    return PdfParser()


@pytest.fixture
def simple_pdf_path(tmp_path: Path) -> Path:
    """Single-page digital PDF with extractable text."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(0, 8, SIMPLE_PDF_TEXT)
    path = tmp_path / "simple.pdf"
    pdf.output(str(path))
    return path


@pytest.fixture
def multi_page_pdf_path(tmp_path: Path) -> Path:
    """Two-page digital PDF — each page has distinct text."""
    from fpdf import FPDF

    pdf = FPDF()
    for page_text in MULTI_PAGE_PDF_TEXTS:
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.multi_cell(0, 8, page_text)
    path = tmp_path / "multi_page.pdf"
    pdf.output(str(path))
    return path


@pytest.fixture
def pdf_with_table_path(tmp_path: Path) -> Path:
    """Single-page PDF with a 3×3 table (bordered cells) preceded by a heading."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    from fpdf.enums import XPos, YPos

    pdf.cell(0, 10, "Personalliste Inf Kp 56/1", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)

    headers = ["Name", "Grad", "Telefon"]
    rows = [
        ["Hans Muster", "Hptm", "079 535 80 46"],
        ["Peter Meier", "Lt", "079 123 45 67"],
    ]
    col_width = 60.0
    row_height = 8.0
    for cell in headers:
        pdf.cell(col_width, row_height, cell, border=1)
    pdf.ln()
    for row in rows:
        for cell in row:
            pdf.cell(col_width, row_height, cell, border=1)
        pdf.ln()

    path = tmp_path / "with_table.pdf"
    pdf.output(str(path))
    return path


@pytest.fixture
def pdf_with_images_path(tmp_path: Path) -> Path:
    """Single-page PDF with 2 embedded raster images and body text."""
    from fpdf import FPDF
    from PIL import Image

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(0, 10, "Dokument mit Bildern")
    pdf.ln(15)

    for _ in range(2):
        img = Image.new("RGB", (100, 100), color=(200, 100, 50))
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        pdf.image(buf, w=40, h=40)
        pdf.ln(5)

    path = tmp_path / "with_images.pdf"
    pdf.output(str(path))
    return path


# ---------------------------------------------------------------------------
# Helper — minimal valid PNG (1×1 red pixel, RGB)
# ---------------------------------------------------------------------------


def _make_minimal_png() -> bytes:
    """Return bytes of a minimal valid 1×1 RGB PNG (red pixel)."""

    def chunk(tag: bytes, data: bytes) -> bytes:
        body = tag + data
        crc = zlib.crc32(body) & 0xFFFFFFFF
        return struct.pack(">I", len(data)) + body + struct.pack(">I", crc)

    signature = b"\x89PNG\r\n\x1a\n"
    ihdr = chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
    # Scanline: filter_byte=0 + R=255, G=0, B=0
    idat = chunk(b"IDAT", zlib.compress(b"\x00\xff\x00\x00"))
    iend = chunk(b"IEND", b"")
    return signature + ihdr + idat + iend


# ---------------------------------------------------------------------------
# XLSX / CSV fixtures
# ---------------------------------------------------------------------------

# PISA 410 column headers (first 8 used in minimal test fixture)
PISA_410_HEADERS = [
    "Versicherten Nr",
    "Einteilung",
    "Dienst bei",
    "Grad",
    "i Gst",
    "Name",
    "Vorname",
    "Funktion",
]

# Matching data row
PISA_410_PERSON_ROW = [
    "756.1234.5678.97",
    "Inf Kp 56/1",
    "Inf Kp 56/1",
    "Hptm",
    "N",
    "Fischer",
    "Lukas",
    "Kdt",
]


@pytest.fixture
def xlsx_csv_parser() -> XlsxCsvParser:
    return XlsxCsvParser()


@pytest.fixture
def simple_xlsx_path(tmp_path: Path) -> Path:
    """Single-sheet XLSX with a header row and two data rows."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Personalliste"
    ws.append(["Name", "Grad", "Einheit"])
    ws.append(["Hans Muster", "Hptm", "Inf Kp 56/1"])
    ws.append(["Peter Meier", "Lt", "Inf Kp 56/2"])
    path = tmp_path / "simple.xlsx"
    wb.save(str(path))
    return path


@pytest.fixture
def multi_sheet_xlsx_path(tmp_path: Path) -> Path:
    """Two-sheet XLSX — 'Offiziere' and 'Unteroffiziere'."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws1 = wb.active
    ws1.title = "Offiziere"
    ws1.append(["Name", "Grad"])
    ws1.append(["Hans Muster", "Hptm"])

    ws2 = wb.create_sheet("Unteroffiziere")
    ws2.append(["Name", "Grad"])
    ws2.append(["Kurt Lehmann", "Sgt"])

    path = tmp_path / "multi_sheet.xlsx"
    wb.save(str(path))
    return path


@pytest.fixture
def pisa_410_xlsx_path(tmp_path: Path) -> Path:
    """PISA 410 XLSX: row 1 = title, row 2 = headers, row 3 = data."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "PISA 410"
    # Row 1: report title in first cell, rest empty (45 cols wide)
    title_row = ["PISA Standardliste 410"] + [""] * (len(PISA_410_HEADERS) - 1)
    ws.append(title_row)
    ws.append(PISA_410_HEADERS)
    ws.append(PISA_410_PERSON_ROW)
    path = tmp_path / "pisa_410.xlsx"
    wb.save(str(path))
    return path


@pytest.fixture
def semicolon_csv_path(tmp_path: Path) -> Path:
    """Semicolon-delimited CSV with header and two data rows."""
    content = (
        "Name;Grad;Einheit\n"
        "Hans Muster;Hptm;Inf Kp 56/1\n"
        "Peter Meier;Lt;Inf Kp 56/2\n"
    )
    path = tmp_path / "semicolon.csv"
    path.write_text(content, encoding="utf-8")
    return path


@pytest.fixture
def comma_csv_path(tmp_path: Path) -> Path:
    """Comma-delimited CSV with header and one data row."""
    content = (
        "Name,Grad,Einheit\n"
        "Peter Meier,Lt,Inf Kp 56/2\n"
    )
    path = tmp_path / "comma.csv"
    path.write_text(content, encoding="utf-8")
    return path


@pytest.fixture
def pisa_410_csv_path(tmp_path: Path) -> Path:
    """PISA 410 CSV: row 1 = title, row 2 = headers, row 3 = data (semicolon)."""
    title = "PISA Standardliste 410" + ";" * (len(PISA_410_HEADERS) - 1)
    header = ";".join(PISA_410_HEADERS)
    data = ";".join(PISA_410_PERSON_ROW)
    content = f"{title}\n{header}\n{data}\n"
    path = tmp_path / "pisa_410.csv"
    path.write_text(content, encoding="utf-8")
    return path
