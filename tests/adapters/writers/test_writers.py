"""Tests for output writers — Markdown, DOCX, CSV, EML."""

from __future__ import annotations

from pathlib import Path

import pytest

from milanon.adapters.writers.csv_writer import CsvWriter
from milanon.adapters.writers.docx_writer import DocxWriter
from milanon.adapters.writers.eml_writer import EmlWriter
from milanon.adapters.writers.markdown_writer import MarkdownWriter
from milanon.domain.entities import AnonymizedDocument, DocumentFormat


def _anon_doc(content: str, structured_content=None) -> AnonymizedDocument:
    return AnonymizedDocument(
        source_path="test.txt",
        output_format=DocumentFormat.MARKDOWN,
        content=content,
        structured_content=structured_content,
    )


class TestMarkdownWriter:
    def test_writes_content_to_file(self, tmp_path: Path):
        doc = _anon_doc("# Anonymisierter Bericht\n\nKein PII mehr.")
        out = tmp_path / "result.md"
        MarkdownWriter().write(doc, out)
        assert out.read_text(encoding="utf-8") == doc.content

    def test_creates_parent_directories(self, tmp_path: Path):
        doc = _anon_doc("text")
        out = tmp_path / "sub" / "dir" / "file.md"
        MarkdownWriter().write(doc, out)
        assert out.exists()

    def test_returns_output_path(self, tmp_path: Path):
        out = tmp_path / "out.md"
        returned = MarkdownWriter().write(_anon_doc("x"), out)
        assert returned == out

    def test_default_extension_is_md(self):
        assert MarkdownWriter().default_extension() == ".md"

    def test_legend_preserved_in_markdown(self, tmp_path: Path):
        content = "<!-- MILANON LEGEND START\n[PERSON_001] = PERSON\nMILANON LEGEND END -->\n\nText."
        out = tmp_path / "out.md"
        MarkdownWriter().write(_anon_doc(content), out)
        assert "MILANON LEGEND" in out.read_text(encoding="utf-8")

    def test_page_breaks_preserved(self, tmp_path: Path):
        content = "Seite 1\n\n---\n\nSeite 2"
        out = tmp_path / "out.md"
        MarkdownWriter().write(_anon_doc(content), out)
        assert "---" in out.read_text(encoding="utf-8")


class TestDocxWriter:
    def test_writes_docx_file(self, tmp_path: Path):
        doc = _anon_doc("[PERSON_001] ist Kommandant.")
        out = tmp_path / "result.docx"
        DocxWriter().write(doc, out)
        assert out.exists()
        assert out.stat().st_size > 0

    def test_content_in_docx(self, tmp_path: Path):
        from docx import Document as DocxDoc

        doc = _anon_doc("[PERSON_001] ist anwesend.")
        out = tmp_path / "result.docx"
        DocxWriter().write(doc, out)
        loaded = DocxDoc(str(out))
        full_text = " ".join(p.text for p in loaded.paragraphs)
        assert "[PERSON_001]" in full_text

    def test_legend_stripped_from_docx(self, tmp_path: Path):
        from docx import Document as DocxDoc

        content = "<!-- MILANON LEGEND START\n[P_001] = PERSON\nMILANON LEGEND END -->\n\nText."
        doc = _anon_doc(content)
        out = tmp_path / "result.docx"
        DocxWriter().write(doc, out)
        loaded = DocxDoc(str(out))
        full_text = " ".join(p.text for p in loaded.paragraphs)
        assert "MILANON LEGEND" not in full_text

    def test_default_extension_is_docx(self):
        assert DocxWriter().default_extension() == ".docx"


class TestCsvWriter:
    def test_writes_csv_file(self, tmp_path: Path):
        doc = _anon_doc("Name | Grad\n[PERSON_001] | Hptm")
        out = tmp_path / "result.csv"
        CsvWriter().write(doc, out)
        assert out.exists()

    def test_structured_content_written_as_csv(self, tmp_path: Path):
        rows = [["Name", "Grad"], ["[PERSON_001]", "Hptm"]]
        doc = _anon_doc("", structured_content={"tables": [rows]})
        out = tmp_path / "result.csv"
        CsvWriter().write(doc, out)
        text = out.read_text(encoding="utf-8")
        assert "[PERSON_001]" in text
        assert "Hptm" in text

    def test_default_extension_is_csv(self):
        assert CsvWriter().default_extension() == ".csv"


class TestEmlWriter:
    def test_writes_eml_file(self, tmp_path: Path):
        doc = _anon_doc("Sehr geehrter [PERSON_001],\n\nIhr Auftrag ist bereit.")
        out = tmp_path / "result.eml"
        EmlWriter().write(doc, out)
        assert out.exists()

    def test_content_preserved_in_eml(self, tmp_path: Path):
        doc = _anon_doc("Body: [PERSON_001]")
        out = tmp_path / "result.eml"
        EmlWriter().write(doc, out)
        assert "[PERSON_001]" in out.read_text(encoding="utf-8")

    def test_legend_stripped_from_eml(self, tmp_path: Path):
        content = "<!-- MILANON LEGEND START\nX\nMILANON LEGEND END -->\n\nBody."
        doc = _anon_doc(content)
        out = tmp_path / "result.eml"
        EmlWriter().write(doc, out)
        assert "MILANON LEGEND" not in out.read_text(encoding="utf-8")

    def test_default_extension_is_eml(self):
        assert EmlWriter().default_extension() == ".eml"
