"""Tests for DocxBefehlWriter — Markdown → DOCX with Armee-Vorlage styles.

Covers BUG-019 (numbered heading recognition), BUG-020 (PDF artifact stripping),
BUG-021 (inline formatting), plus existing functionality.
"""

from pathlib import Path

import pytest
from docx import Document

from milanon.adapters.writers.docx_befehl_writer import (
    DocxBefehlWriter,
    _strip_pdf_artifacts,
)

TEMPLATE_PATH = Path("data/templates/docx/befehl_vorlage.docx")


@pytest.fixture
def writer():
    return DocxBefehlWriter()


@pytest.fixture
def output_path(tmp_path):
    return tmp_path / "output.docx"


class TestDocxBefehlWriter:
    def test_markdown_to_docx_basic_structure(self, writer, output_path):
        """Basic Markdown produces a valid DOCX with correct number of paragraphs."""
        md = (
            "# Grundlagen\n"
            "- Lage X;\n"
            "- Lage Y.\n"
            "### 1 Orientierung\n"
            "#### 1.1 Bedrohung\n"
            "Text body here.\n"
        )
        result = writer.write(md, TEMPLATE_PATH, output_path)

        assert result == output_path
        assert output_path.exists()

        doc = Document(str(output_path))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "Grundlagen" in texts
        assert "Lage X;" in texts
        assert "1 Orientierung" in texts
        assert "1.1 Bedrohung" in texts
        assert "Text body here." in texts

    def test_heading_level_detection(self, writer, output_path):
        """Each Markdown heading level maps to the correct DOCX style."""
        md = (
            "# Grundlagen\n"
            "## DECKNAME\n"
            "### 1 Orientierung\n"
            "#### 1.1 Bedrohung\n"
            "##### 1.1.1 Bestimmend\n"
        )
        writer.write(md, TEMPLATE_PATH, output_path)
        doc = Document(str(output_path))

        style_map = {
            p.text: p.style.name for p in doc.paragraphs if p.text.strip()
        }
        assert style_map["Grundlagen"] == "Heading 1"
        assert style_map["DECKNAME"] == "Subject heading"
        assert style_map["1 Orientierung"] == "1. Main title"
        assert style_map["1.1 Bedrohung"] == "1.1 Title"
        assert style_map["1.1.1 Bestimmend"] == "1.1.1 Title"

    def test_table_detection_and_conversion(self, writer, output_path):
        """Markdown pipe tables are converted to DOCX tables with 2 columns."""
        md = (
            "### 3 Aufträge\n"
            "| Element | Auftrag |\n"
            "|---|---|\n"
            "| Z Ambos | NIMMT und SICHERT Ags WEST |\n"
            "| Z Canale | SPERRT auf Höhe CHUR |\n"
        )
        writer.write(md, TEMPLATE_PATH, output_path)
        doc = Document(str(output_path))

        assert len(doc.tables) >= 1
        table = doc.tables[0]
        assert len(table.rows) == 2
        assert len(table.columns) == 2
        assert "Z Ambos" in table.rows[0].cells[0].text
        assert "Z Canale" in table.rows[1].cells[0].text


class TestBug019NumberedHeadings:
    """BUG-019: Detect numbered heading patterns from PDF extraction."""

    def test_main_section_heading(self, writer, output_path):
        """'1. Orientierung' is detected as a main section heading."""
        md = "1. Orientierung\n"
        writer.write(md, TEMPLATE_PATH, output_path)
        doc = Document(str(output_path))

        paras = [p for p in doc.paragraphs if p.text.strip()]
        assert len(paras) == 1
        assert paras[0].text == "1. Orientierung"
        assert paras[0].style.name == "1. Main title"

    def test_subsection_heading(self, writer, output_path):
        """'1.1. Grundlagen' is detected as a subsection heading."""
        md = "1.1. Grundlagen\n"
        writer.write(md, TEMPLATE_PATH, output_path)
        doc = Document(str(output_path))

        paras = [p for p in doc.paragraphs if p.text.strip()]
        assert len(paras) == 1
        assert paras[0].text == "1.1. Grundlagen"
        assert paras[0].style.name == "1.1 Title"

    def test_subsection_heading_no_trailing_dot(self, writer, output_path):
        """'3.2 Kdt Stabskp' (no trailing dot) is also a subsection heading."""
        md = "3.2 Kdt Stabskp\n"
        writer.write(md, TEMPLATE_PATH, output_path)
        doc = Document(str(output_path))

        paras = [p for p in doc.paragraphs if p.text.strip()]
        assert len(paras) == 1
        assert paras[0].style.name == "1.1 Title"

    def test_sub_subsection_heading(self, writer, output_path):
        """'1.2.1. Erhaltener Auftrag' is detected as a sub-subsection."""
        md = "1.2.1. Erhaltener Auftrag\n"
        writer.write(md, TEMPLATE_PATH, output_path)
        doc = Document(str(output_path))

        paras = [p for p in doc.paragraphs if p.text.strip()]
        assert len(paras) == 1
        assert paras[0].text == "1.2.1. Erhaltener Auftrag"
        assert paras[0].style.name == "1.1.1 Title"

    def test_sub_subsection_no_trailing_dot(self, writer, output_path):
        """'4.3.5 Bedrohungsstufe' (no trailing dot) is a sub-subsection."""
        md = "4.3.5 Bedrohungsstufe\n"
        writer.write(md, TEMPLATE_PATH, output_path)
        doc = Document(str(output_path))

        paras = [p for p in doc.paragraphs if p.text.strip()]
        assert len(paras) == 1
        assert paras[0].style.name == "1.1.1 Title"

    def test_numbered_heading_priority_over_list(self, writer, output_path):
        """Main section headings take priority over ordered list detection."""
        md = (
            "1. Orientierung\n"
            "2. Absicht\n"
            "5. Standorte\n"
        )
        writer.write(md, TEMPLATE_PATH, output_path)
        doc = Document(str(output_path))

        paras = [p for p in doc.paragraphs if p.text.strip()]
        assert len(paras) == 3
        for p in paras:
            assert p.style.name == "1. Main title"

    def test_lowercase_list_item_not_heading(self, writer, output_path):
        """'1. item text' with lowercase is treated as ordered list."""
        md = "1. item text in a list\n"
        writer.write(md, TEMPLATE_PATH, output_path)
        doc = Document(str(output_path))

        paras = [p for p in doc.paragraphs if p.text.strip()]
        assert len(paras) == 1
        assert paras[0].style.name == "Numbered List 2"

    def test_full_document_with_numbered_headings(self, writer, output_path):
        """Complete document with numbered headings renders correctly."""
        md = (
            "1. Orientierung\n"
            "1.1. Grundlagen\n"
            "- Lage gemäss Bat Bf;\n"
            "1.2. Feind\n"
            "1.2.1. Erhaltener Auftrag\n"
            "Feind greift an.\n"
            "2. Absicht\n"
            "Kp verteidigt.\n"
        )
        writer.write(md, TEMPLATE_PATH, output_path)
        doc = Document(str(output_path))

        style_map = {
            p.text: p.style.name for p in doc.paragraphs if p.text.strip()
        }
        assert style_map["1. Orientierung"] == "1. Main title"
        assert style_map["1.1. Grundlagen"] == "1.1 Title"
        assert style_map["Lage gemäss Bat Bf;"] == "Bullet List 1"
        assert style_map["1.2. Feind"] == "1.1 Title"
        assert style_map["1.2.1. Erhaltener Auftrag"] == "1.1.1 Title"
        assert style_map["Feind greift an."] == "Text Indent"
        assert style_map["2. Absicht"] == "1. Main title"
        assert style_map["Kp verteidigt."] == "Text Indent"


class TestBug020PdfArtifacts:
    """BUG-020: Strip PDF page artifacts from input."""

    def test_strip_docx_filename(self):
        """Lines ending with .docx are removed."""
        lines = [
            "1. Orientierung",
            "WK25_InfBat56_1.00_Kdt_Allgemeiner_Bf_WK25.docx",
            "Some normal text.",
        ]
        result = _strip_pdf_artifacts(lines)
        assert len(result) == 2
        assert "Some normal text." in result

    def test_strip_page_number_only(self):
        """Lines that are only page numbers like '1/7' are removed."""
        lines = ["1/7", "2/7", "Normal text"]
        result = _strip_pdf_artifacts(lines)
        assert result == ["Normal text"]

    def test_strip_footer_with_quoted_string(self):
        """Footer lines with quoted text and page number are removed."""
        lines = [
            'Kdt Inf Bat 56, "WK25", Allgemeiner Befehl 2/7',
            "Normal paragraph text.",
        ]
        result = _strip_pdf_artifacts(lines)
        assert len(result) == 1
        assert result[0] == "Normal paragraph text."

    def test_preserve_normal_lines(self):
        """Normal content lines are preserved."""
        lines = [
            "1. Orientierung",
            "- Lage gemäss Bat Bf;",
            "",
            "Body text here.",
        ]
        result = _strip_pdf_artifacts(lines)
        assert len(result) == 4

    def test_preserve_empty_lines(self):
        """Empty lines are preserved (spacing)."""
        lines = ["", "text", ""]
        result = _strip_pdf_artifacts(lines)
        assert len(result) == 3

    def test_integration_artifacts_stripped_in_write(
        self, writer, output_path
    ):
        """PDF artifacts are stripped during write() — not present in output."""
        md = (
            "WK25_InfBat56_1.00_Kdt_Allgemeiner_Bf_WK25.docx\n"
            "1/7\n"
            "# Grundlagen\n"
            '  Kdt Inf Bat 56, "WK25", Allgemeiner Befehl 2/7\n'
            "Normal text.\n"
        )
        writer.write(md, TEMPLATE_PATH, output_path)
        doc = Document(str(output_path))

        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "Grundlagen" in texts
        assert "Normal text." in texts
        # Artifacts should NOT appear
        assert not any(".docx" in t for t in texts)
        assert "1/7" not in texts


class TestBug021InlineFormatting:
    """BUG-021: Bold/italic inline formatting in DOCX output."""

    def test_bold_text(self, writer, output_path):
        """**bold** text renders with bold runs."""
        md = "**Wichtig** ist das.\n"
        writer.write(md, TEMPLATE_PATH, output_path)
        doc = Document(str(output_path))

        paras = [p for p in doc.paragraphs if p.text.strip()]
        assert len(paras) == 1
        runs = paras[0].runs
        assert any(r.bold and r.text == "Wichtig" for r in runs)
        assert any(not r.bold and "ist das." in r.text for r in runs)

    def test_italic_text(self, writer, output_path):
        """*italic* text renders with italic runs."""
        md = "Das ist *kursiv* hier.\n"
        writer.write(md, TEMPLATE_PATH, output_path)
        doc = Document(str(output_path))

        paras = [p for p in doc.paragraphs if p.text.strip()]
        assert len(paras) == 1
        runs = paras[0].runs
        assert any(r.italic and r.text == "kursiv" for r in runs)

    def test_bold_italic_combined(self, writer, output_path):
        """***bold+italic*** renders with both bold and italic."""
        md = "***Achtung*** Feind.\n"
        writer.write(md, TEMPLATE_PATH, output_path)
        doc = Document(str(output_path))

        paras = [p for p in doc.paragraphs if p.text.strip()]
        runs = paras[0].runs
        assert any(
            r.bold and r.italic and r.text == "Achtung" for r in runs
        )

    def test_inline_formatting_in_bullet(self, writer, output_path):
        """Inline formatting works inside bullet list items."""
        md = "- **Lage** gemäss Bat Bf;\n"
        writer.write(md, TEMPLATE_PATH, output_path)
        doc = Document(str(output_path))

        paras = [p for p in doc.paragraphs if p.text.strip()]
        assert len(paras) == 1
        assert paras[0].style.name == "Bullet List 1"
        runs = paras[0].runs
        assert any(r.bold and r.text == "Lage" for r in runs)

    def test_inline_formatting_in_heading(self, writer, output_path):
        """Inline formatting works inside headings."""
        md = "### 1 **Orientierung**\n"
        writer.write(md, TEMPLATE_PATH, output_path)
        doc = Document(str(output_path))

        paras = [p for p in doc.paragraphs if p.text.strip()]
        assert paras[0].style.name == "1. Main title"
        runs = paras[0].runs
        assert any(r.bold and r.text == "Orientierung" for r in runs)

    def test_table_cells_have_inline_formatting(self, writer, output_path):
        """Table cell content supports inline formatting."""
        md = (
            "| Element | Auftrag |\n"
            "|---|---|\n"
            "| **Z Alpha** | NIMMT Ags |\n"
        )
        writer.write(md, TEMPLATE_PATH, output_path)
        doc = Document(str(output_path))

        assert len(doc.tables) >= 1
        left_cell = doc.tables[0].rows[0].cells[0]
        runs = left_cell.paragraphs[0].runs
        assert any(r.bold and r.text == "Z Alpha" for r in runs)


class TestHorizontalRule:
    """Horizontal rules (---) become page breaks."""

    def test_hr_creates_page_break(self, writer, output_path):
        """'---' line produces a page break element."""
        md = "Text before.\n---\nText after.\n"
        writer.write(md, TEMPLATE_PATH, output_path)
        doc = Document(str(output_path))

        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "Text before." in texts
        assert "Text after." in texts

        # Check that a w:br element with type=page exists
        from docx.oxml.ns import qn as _qn

        breaks = doc.element.body.findall(
            f".//{_qn('w:br')}[@{_qn('w:type')}='page']"
        )
        assert len(breaks) >= 1


class TestBlockquote:
    """Blockquotes (> text) render as italic Text Indent."""

    def test_blockquote_renders_italic(self, writer, output_path):
        """> text renders as italic Text Indent paragraph."""
        md = "> Dies ist ein Zitat.\n"
        writer.write(md, TEMPLATE_PATH, output_path)
        doc = Document(str(output_path))

        paras = [p for p in doc.paragraphs if p.text.strip()]
        assert len(paras) == 1
        assert paras[0].style.name == "Text Indent"
        assert paras[0].runs[0].italic


class TestHtmlCommentStripping:
    """HTML comments are stripped from output."""

    def test_html_comments_removed(self, writer, output_path):
        """<!-- comments --> are stripped before rendering."""
        md = "Text <!-- this is a comment --> here.\n"
        writer.write(md, TEMPLATE_PATH, output_path)
        doc = Document(str(output_path))

        paras = [p for p in doc.paragraphs if p.text.strip()]
        assert len(paras) == 1
        assert "comment" not in paras[0].text
        assert "Text  here." in paras[0].text

    def test_multiline_html_comment_stripped(self, writer, output_path):
        """Multi-line HTML comments are also stripped."""
        md = "Before.\n<!-- multi\nline\ncomment -->\nAfter.\n"
        writer.write(md, TEMPLATE_PATH, output_path)
        doc = Document(str(output_path))

        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "Before." in texts
        assert "After." in texts
        assert not any("comment" in t for t in texts)


class TestLegendStripping:
    """Legend block from anonymizer is stripped before rendering."""

    def test_legend_block_not_in_output(self, writer, output_path):
        """MILANON LEGEND HTML comment block is stripped from output."""
        md = (
            "# Grundlagen\n"
            "Text.\n"
            "\n"
            "<!-- MILANON LEGEND START\n"
            "[PERSON_001] = Hptm Muster\n"
            "MILANON LEGEND END -->\n"
        )
        writer.write(md, TEMPLATE_PATH, output_path)
        doc = Document(str(output_path))

        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert "Grundlagen" in texts
        assert "Text." in texts
        assert not any("PERSON_001" in t for t in texts)
        assert not any("LEGEND" in t for t in texts)
