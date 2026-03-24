"""Tests for ExportDocxUseCase — DOCX export with optional de-anonymization."""

from pathlib import Path
from unittest.mock import MagicMock

import pytest
from docx import Document

from milanon.adapters.writers.docx_befehl_writer import DocxBefehlWriter
from milanon.domain.entities import EntityMapping, EntityType
from milanon.usecases.export_docx import ExportDocxUseCase

TEMPLATE_PATH = Path("data/templates/docx/befehl_vorlage.docx")


def _make_mapping(entity_type: EntityType, original: str, placeholder: str):
    return EntityMapping(
        entity_type=entity_type,
        original_value=original,
        placeholder=placeholder,
    )


@pytest.fixture
def repo():
    return MagicMock()


@pytest.fixture
def writer():
    return DocxBefehlWriter()


@pytest.fixture
def use_case(repo, writer):
    return ExportDocxUseCase(repository=repo, writer=writer)


@pytest.fixture
def md_file(tmp_path):
    """Create a sample Markdown Befehl file."""
    path = tmp_path / "befehl.md"
    path.write_text(
        "# Grundlagen\n"
        "- Lage gemäss Bat Bf;\n"
        "- Bedrohung aktuell.\n"
        "### 1 Orientierung\n"
        "#### 1.1 Bedrohung\n"
        "Der Gegner greift an.\n"
        "### 3 Aufträge\n"
        "| Element | Auftrag |\n"
        "|---|---|\n"
        "| Z Ambos | NIMMT und SICHERT Ags WEST |\n"
        "| Z Canale | SPERRT auf Höhe CHUR |\n",
        encoding="utf-8",
    )
    return path


@pytest.fixture
def md_file_with_placeholders(tmp_path):
    """Markdown file containing anonymized placeholders."""
    path = tmp_path / "befehl_anon.md"
    path.write_text(
        "# Grundlagen\n"
        "### 1 Orientierung\n"
        "[PERSON_001] führt den Angriff.\n"
        "[ORT_003] ist das Ziel.\n"
        "### 3 Aufträge\n"
        "| Element | Auftrag |\n"
        "|---|---|\n"
        "| [EINHEIT_001] | NIMMT [ORT_003] |\n",
        encoding="utf-8",
    )
    return path


class TestExportDocxUseCase:
    def test_export_creates_docx_file(self, use_case, md_file, tmp_path):
        output = tmp_path / "output.docx"
        result = use_case.execute(md_file, output, TEMPLATE_PATH)

        assert result == output
        assert output.exists()
        assert output.stat().st_size > 0

    def test_heading_styles_applied_correctly(self, use_case, md_file, tmp_path):
        output = tmp_path / "output.docx"
        use_case.execute(md_file, output, TEMPLATE_PATH)

        doc = Document(str(output))
        style_map = {p.text: p.style.name for p in doc.paragraphs if p.text.strip()}

        assert style_map.get("Grundlagen") == "Heading 1"
        assert style_map.get("1 Orientierung") == "1. Main title"
        assert style_map.get("1.1 Bedrohung") == "1.1 Title"

    def test_auftraege_table_created_with_correct_columns(
        self, use_case, md_file, tmp_path
    ):
        output = tmp_path / "output.docx"
        use_case.execute(md_file, output, TEMPLATE_PATH)

        doc = Document(str(output))
        assert len(doc.tables) >= 1
        table = doc.tables[0]
        assert len(table.columns) == 2
        assert "Z Ambos" in table.rows[0].cells[0].text

    def test_bullet_lists_use_correct_style(self, use_case, md_file, tmp_path):
        output = tmp_path / "output.docx"
        use_case.execute(md_file, output, TEMPLATE_PATH)

        doc = Document(str(output))
        bullet_paras = [
            p for p in doc.paragraphs if p.style.name == "Bullet List 1"
        ]
        bullet_texts = [p.text for p in bullet_paras]
        assert "Lage gemäss Bat Bf;" in bullet_texts
        assert "Bedrohung aktuell." in bullet_texts

    def test_deanonymize_replaces_person_placeholder(
        self, repo, writer, md_file_with_placeholders, tmp_path
    ):
        repo.get_placeholder.side_effect = lambda ph: {
            "[PERSON_001]": _make_mapping(
                EntityType.PERSON, "Hptm Degen", "[PERSON_001]"
            ),
            "[ORT_003]": _make_mapping(EntityType.ORT, "Bern", "[ORT_003]"),
            "[EINHEIT_001]": _make_mapping(
                EntityType.EINHEIT, "Inf Kp 56/1", "[EINHEIT_001]"
            ),
        }.get(ph)

        uc = ExportDocxUseCase(repository=repo, writer=writer)
        output = tmp_path / "output.docx"
        uc.execute(md_file_with_placeholders, output, TEMPLATE_PATH, deanonymize=True)

        doc = Document(str(output))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Hptm Degen" in all_text
        assert "[PERSON_001]" not in all_text

    def test_deanonymize_replaces_ort_placeholder(
        self, repo, writer, md_file_with_placeholders, tmp_path
    ):
        repo.get_placeholder.side_effect = lambda ph: {
            "[PERSON_001]": _make_mapping(
                EntityType.PERSON, "Hptm Degen", "[PERSON_001]"
            ),
            "[ORT_003]": _make_mapping(EntityType.ORT, "Bern", "[ORT_003]"),
            "[EINHEIT_001]": _make_mapping(
                EntityType.EINHEIT, "Inf Kp 56/1", "[EINHEIT_001]"
            ),
        }.get(ph)

        uc = ExportDocxUseCase(repository=repo, writer=writer)
        output = tmp_path / "output.docx"
        uc.execute(md_file_with_placeholders, output, TEMPLATE_PATH, deanonymize=True)

        doc = Document(str(output))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Bern" in all_text
        assert "[ORT_003]" not in all_text

    def test_deanonymize_in_table_cells(
        self, repo, writer, md_file_with_placeholders, tmp_path
    ):
        repo.get_placeholder.side_effect = lambda ph: {
            "[PERSON_001]": _make_mapping(
                EntityType.PERSON, "Hptm Degen", "[PERSON_001]"
            ),
            "[ORT_003]": _make_mapping(EntityType.ORT, "Bern", "[ORT_003]"),
            "[EINHEIT_001]": _make_mapping(
                EntityType.EINHEIT, "Inf Kp 56/1", "[EINHEIT_001]"
            ),
        }.get(ph)

        uc = ExportDocxUseCase(repository=repo, writer=writer)
        output = tmp_path / "output.docx"
        uc.execute(md_file_with_placeholders, output, TEMPLATE_PATH, deanonymize=True)

        doc = Document(str(output))
        table_text = " ".join(
            cell.text
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        )
        assert "Inf Kp 56/1" in table_text
        assert "Bern" in table_text
        assert "[EINHEIT_001]" not in table_text
        assert "[ORT_003]" not in table_text

    def test_template_header_footer_preserved(self, use_case, md_file, tmp_path):
        output = tmp_path / "output.docx"
        use_case.execute(md_file, output, TEMPLATE_PATH)

        doc = Document(str(output))
        # Template has 1 section with header and footer
        assert len(doc.sections) >= 1
        section = doc.sections[0]
        assert len(section.header.paragraphs) >= 1
        assert len(section.footer.paragraphs) >= 1
